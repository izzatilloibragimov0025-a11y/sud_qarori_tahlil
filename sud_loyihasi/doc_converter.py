"""
doc_converter.py — Universal hujjat → matn converter
Qo'llab-quvvatlanadi: .pdf, .doc, .docx

Strategiya: format avtomatik aniqlanadi (file signature orqali),
keyin mos extractor ishlatiladi. Har format uchun bir nechta fallback
mavjud — bittasi fail bo'lsa, keyingisi sinaladi.

ISHLATISH:
    from doc_converter import doc_to_text, batch_convert
    
    # Bitta fayl
    text = doc_to_text("sud_fayllari/ID_1476730.doc")
    
    # Butun papka
    batch_convert("sud_fayllari/", "matnlar/")
"""
import os
import subprocess
import json
from pathlib import Path
from typing import Optional


# ============================================================
# FORMAT ANIQLASH (file signature orqali)
# ============================================================
def detect_format(filepath: str) -> str:
    """Faylning haqiqiy formatini aniqlash (kengaytmaga emas, magic byte'ga qarab)"""
    try:
        with open(filepath, "rb") as f:
            head = f.read(8)
    except Exception:
        return "unknown"

    if head[:4] == b"%PDF":
        return "pdf"
    if head[:4] == b"\xd0\xcf\x11\xe0":  # OLE2 — klassik .doc/.xls
        return "doc"
    if head[:2] == b"PK":  # ZIP — .docx, .xlsx
        return "docx"
    return "unknown"


# ============================================================
# EXTRACTORS — har format uchun
# ============================================================
def extract_pdf(filepath: str) -> Optional[str]:
    """PDF dan matn olish (pdfplumber)"""
    try:
        import pdfplumber
        with pdfplumber.open(filepath) as pdf:
            pages = [p.extract_text() or "" for p in pdf.pages]
        text = "\n".join(pages).strip()
        return text if text else None
    except Exception as e:
        print(f"   pdfplumber xato: {e}")
        return None


def extract_docx(filepath: str) -> Optional[str]:
    """DOCX dan matn olish (python-docx)"""
    try:
        import docx
        doc = docx.Document(filepath)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Jadvallardagi matnni ham olish
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        text = "\n".join(paragraphs).strip()
        return text if text else None
    except Exception as e:
        print(f"   python-docx xato: {e}")
        return None


def extract_doc_libreoffice(filepath: str, tmp_dir: str = "tmp_convert") -> Optional[str]:
    """LibreOffice headless orqali .doc → .txt
    
    Bu eng ishonchli usul, lekin LibreOffice o'rnatilgan bo'lishi shart.
    Windows: https://www.libreoffice.org/download/
    """
    Path(tmp_dir).mkdir(exist_ok=True)
    
    # LibreOffice ning turli nomlari (Windows va Linux)
    soffice_names = ["soffice", "libreoffice", 
                     r"C:\Program Files\LibreOffice\program\soffice.exe",
                     r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"]
    
    for soffice in soffice_names:
        try:
            cmd = [soffice, "--headless", "--convert-to", "txt:Text (encoded):UTF8",
                   "--outdir", tmp_dir, str(Path(filepath).resolve())]
            result = subprocess.run(cmd, capture_output=True, timeout=60)
            
            if result.returncode == 0:
                txt_path = Path(tmp_dir) / (Path(filepath).stem + ".txt")
                if txt_path.exists():
                    text = txt_path.read_text(encoding="utf-8", errors="replace")
                    txt_path.unlink()  # cleanup
                    return text.strip() if text.strip() else None
        except (FileNotFoundError, subprocess.TimeoutExpired):
            continue
        except Exception as e:
            print(f"   LibreOffice xato ({soffice}): {e}")
            continue
    return None


def extract_doc_word(filepath: str) -> Optional[str]:
    """Microsoft Word orqali (faqat Windows + Word o'rnatilgan)"""
    try:
        import win32com.client
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        try:
            doc = word.Documents.Open(str(Path(filepath).resolve()), ReadOnly=True)
            text = doc.Content.Text
            doc.Close(SaveChanges=False)
            return text.strip() if text.strip() else None
        finally:
            word.Quit()
    except ImportError:
        return None  # pywin32 o'rnatilmagan
    except Exception as e:
        print(f"   Word COM xato: {e}")
        return None


def extract_doc_olefile(filepath: str) -> Optional[str]:
    """Oxirgi chora — olefile bilan binary parsing
    
    Eski Word fayllaridan asosiy matnni topishga harakat qiladi.
    Sifati past, lekin LibreOffice/Word yo'q bo'lganda yaqin keladi.
    """
    try:
        import olefile
        if not olefile.isOleFile(filepath):
            return None
        
        ole = olefile.OleFileIO(filepath)
        try:
            if not ole.exists("WordDocument"):
                return None
            stream = ole.openstream("WordDocument")
            data = stream.read()
        finally:
            ole.close()
        
        # Encoding'larni navbat bilan sinab ko'rish
        for encoding in ["cp1251", "utf-8", "cp1252", "latin-1"]:
            try:
                text = data.decode(encoding, errors="ignore")
                # Faqat printable belgilarni qoldirish
                cleaned = "".join(c for c in text if c.isprintable() or c in "\n\r\t ")
                # Juda ko'p shovqin bo'lsa, juda ham past sifat — keyingisini sinaymiz
                if len(cleaned.strip()) > 200:
                    return cleaned
            except Exception:
                continue
        return None
    except ImportError:
        return None
    except Exception as e:
        print(f"   olefile xato: {e}")
        return None


# ============================================================
# ASOSIY API
# ============================================================
def doc_to_text(filepath: str, verbose: bool = True) -> Optional[str]:
    """Faylni matnga aylantirish. Format avtomatik aniqlanadi.
    
    Returns:
        str: muvaffaqiyatli ekstraksiya qilingan matn
        None: hech qanday usul ishlamadi
    """
    if not os.path.exists(filepath):
        if verbose:
            print(f"   Fayl topilmadi: {filepath}")
        return None
    
    fmt = detect_format(filepath)
    if verbose:
        print(f"   Format: {fmt}")
    
    if fmt == "pdf":
        return extract_pdf(filepath)
    
    if fmt == "docx":
        return extract_docx(filepath)
    
    if fmt == "doc":
        # Strategiyalarni navbat bilan sinash
        for name, extractor in [
            ("LibreOffice", lambda: extract_doc_libreoffice(filepath)),
            ("MS Word",     lambda: extract_doc_word(filepath)),
            ("olefile",     lambda: extract_doc_olefile(filepath)),
        ]:
            if verbose:
                print(f"   Sinab ko'rilmoqda: {name}...")
            text = extractor()
            if text and len(text.strip()) > 100:
                if verbose:
                    print(f"   ✓ {name} muvaffaqiyatli ({len(text)} belgi)")
                return text
        return None
    
    if verbose:
        print(f"   Noma'lum format: {fmt}")
    return None


# ============================================================
# BATCH (butun papkani konvertatsiya qilish)
# ============================================================
def batch_convert(input_dir: str, output_dir: str = "matnlar",
                  manifest_path: str = "conversion_manifest.json") -> dict:
    """Butun papkadagi fayllarni matnga o'tkazish.
    
    Args:
        input_dir: kirish papkasi (.doc/.pdf fayllar)
        output_dir: chiqish papkasi (.txt fayllar)
        manifest_path: hisobot fayli (qaysi muvaffaqiyatli/qaysi yo'q)
    
    Returns:
        dict: {"success": [...], "failed": [...], "stats": {...}}
    """
    in_dir = Path(input_dir)
    out_dir = Path(output_dir)
    out_dir.mkdir(exist_ok=True)
    
    if not in_dir.exists():
        print(f"❌ Kirish papkasi topilmadi: {in_dir}")
        return {"success": [], "failed": [], "stats": {}}
    
    files = list(in_dir.glob("*.doc")) + list(in_dir.glob("*.pdf")) + \
            list(in_dir.glob("*.docx"))
    
    print(f"\n{'='*70}")
    print(f"  Batch konvertatsiya: {len(files)} ta fayl")
    print(f"  Kirish:  {in_dir.absolute()}")
    print(f"  Chiqish: {out_dir.absolute()}")
    print(f"{'='*70}\n")
    
    success, failed = [], []
    
    for i, filepath in enumerate(files, 1):
        print(f"[{i}/{len(files)}] {filepath.name}")
        text = doc_to_text(str(filepath), verbose=True)
        
        if text:
            out_path = out_dir / (filepath.stem + ".txt")
            out_path.write_text(text, encoding="utf-8")
            success.append({"file": filepath.name, "size": len(text), "out": out_path.name})
            print(f"   → saqlandi: {out_path.name}")
        else:
            failed.append({"file": filepath.name, "reason": "extraction failed"})
            print(f"   ✗ matn ajratib bo'lmadi")
        print()
    
    # Hisobot
    manifest = {
        "success": success,
        "failed": failed,
        "stats": {
            "total": len(files),
            "succeeded": len(success),
            "failed": len(failed),
            "success_rate": f"{len(success)/max(1,len(files))*100:.1f}%",
        },
    }
    
    with open(manifest_path, "w", encoding="utf-8") as f:
        json.dump(manifest, f, ensure_ascii=False, indent=2)
    
    print(f"{'='*70}")
    print(f"  YAKUN: {len(success)}/{len(files)} muvaffaqiyatli "
          f"({len(success)/max(1,len(files))*100:.1f}%)")
    print(f"  Hisobot: {manifest_path}")
    print(f"{'='*70}\n")
    
    if failed:
        print("Buzilgan fayllar:")
        for item in failed:
            print(f"   ✗ {item['file']}")
    
    return manifest


# ============================================================
# COMMAND LINE
# ============================================================
if __name__ == "__main__":
    import sys
    
    if len(sys.argv) < 2:
        # Default: sud_fayllari → matnlar
        batch_convert("sud_fayllari", "matnlar")
    elif len(sys.argv) == 2:
        # Bitta fayl
        text = doc_to_text(sys.argv[1])
        if text:
            print(text)
        else:
            sys.exit(1)
    else:
        # Custom papkalar
        batch_convert(sys.argv[1], sys.argv[2])
